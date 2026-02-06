import re
import shutil
from pathlib import Path
from datetime import datetime
from typing import Optional, Union, List, Tuple, Dict
import logging
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# 设置日志
logger = logging.getLogger(__name__)

def is_file_within_workspace(file_path: Path, workspace_path: Path) -> bool:
    """检查文件是否在工作区内"""
    try:
        return str(file_path.resolve()).startswith(str(workspace_path.resolve()))
    except Exception:
        return False


def read_word_file(file_path: Union[str, Path],
                   workspace_path: Optional[Union[str, Path]] = None) -> Tuple[bool, Dict]:
    """读取Word文档内容

    Args:
        file_path: Word文档路径
        workspace_path: 工作区路径

    Returns:
        Tuple[bool, Dict]: (是否成功, 包含文本、段落和表格信息)
    """
    try:
        # 处理文件路径
        if isinstance(file_path, str):
            file_path = Path(file_path)

        # 处理工作区路径
        if workspace_path is not None:
            workspace = Path(workspace_path)
        else:
            workspace = Path.cwd()

        # 处理相对路径
        if not file_path.is_absolute():
            file_path = workspace / file_path

        # 安全检查
        if not is_file_within_workspace(file_path, workspace):
            return False, {"error": "文件不在当前工作区内"}

        if not file_path.exists():
            return False, {"error": f"文件不存在: {file_path}"}

        if file_path.suffix.lower() not in ['.docx', '.doc']:
            return False, {"error": "文件不是Word文档格式"}

        # 读取Word文档
        doc = Document(file_path)

        # 提取文本内容
        full_text = []
        paragraphs = []

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                full_text.append(paragraph.text)
                paragraphs.append({
                    "index": i,
                    "text": paragraph.text,
                    "style": paragraph.style.name
                })

        # 提取表格内容
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)

            tables.append({
                "index": i,
                "data": table_data,
                "rows": len(table.rows),
                "columns": len(table.columns)
            })

        result = {
            "file_path": str(file_path),
            "full_text": "\n".join(full_text),
            "paragraphs": paragraphs,
            "tables": tables,
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables)
        }

        logger.info(f"成功读取Word文档: {file_path}")
        return True, result

    except Exception as e:
        logger.error(f"读取Word文档失败: {e}")
        return False, {"error": f"读取Word文档失败: {e}"}


def create_word_file(file_path: Union[str, Path],
                     content: List[str] = None,
                     workspace_path: Optional[Union[str, Path]] = None,
                     overwrite: bool = False) -> Tuple[bool, str]:
    """创建新的Word文档

    Args:
        file_path: 文件路径
        content: 段落内容列表
        workspace_path: 工作区路径
        overwrite: 是否覆盖已存在文件

    Returns:
        Tuple[bool, str]: 是否成功和消息
    """
    try:
        # 处理文件路径
        if isinstance(file_path, str):
            file_path = Path(file_path)

        # 处理工作区路径
        if workspace_path is not None:
            workspace = Path(workspace_path)
        else:
            workspace = Path.cwd()

        # 处理相对路径
        if not file_path.is_absolute():
            file_path = workspace / file_path

        # 确保文件在工作区内
        if not is_file_within_workspace(file_path, workspace):
            return False, "错误：文件路径不在当前工作区内"

        # 确保文件扩展名为.docx
        if file_path.suffix.lower() != ".docx":
            file_path = file_path.with_suffix(".docx")

        # 检查文件是否已存在
        if file_path.exists() and not overwrite:
            return False, f"错误：文件已存在: {file_path}。如需覆盖请设置 overwrite=True"

        # 创建父目录（如果不存在）
        file_path.parent.mkdir(parents=True, exist_ok=True)

        # 创建Word文档
        doc = Document()

        # 添加内容
        if content:
            for para_text in content:
                doc.add_paragraph(para_text)
        else:
            doc.add_paragraph("新建Word文档")

        # 保存文档
        doc.save(file_path)

        logger.info(f"成功创建Word文档: {file_path}")
        return True, f"Word文档创建成功: {file_path}"

    except Exception as e:
        return False, f"创建Word文档失败: {e}"


def add_paragraph_to_word(file_path: Union[str, Path],
                          paragraph_text: str,
                          workspace_path: Optional[Union[str, Path]] = None) -> Tuple[bool, str]:
    """向Word文档添加段落

    Args:
        file_path: Word文档路径
        paragraph_text: 要添加的段落文本
        workspace_path: 工作区路径

    Returns:
        Tuple[bool, str]: 是否成功和消息
    """
    try:
        # 处理文件路径
        if isinstance(file_path, str):
            file_path = Path(file_path)

        # 处理工作区路径
        if workspace_path is not None:
            workspace = Path(workspace_path)
        else:
            workspace = Path.cwd()

        # 处理相对路径
        if not file_path.is_absolute():
            file_path = workspace / file_path

        # 安全检查
        if not is_file_within_workspace(file_path, workspace):
            return False, "错误：文件不在当前工作区内"

        if not file_path.exists():
            return False, f"错误：文件不存在: {file_path}"

        # 打开文档并添加段落
        doc = Document(file_path)
        doc.add_paragraph(paragraph_text)
        doc.save(file_path)

        logger.info(f"成功向Word文档添加段落: {file_path}")
        return True, f"成功添加段落到文档"

    except Exception as e:
        return False, f"添加段落失败: {e}"


def add_table_to_word(file_path: Union[str, Path],
                      table_data: List[List[str]],
                      workspace_path: Optional[Union[str, Path]] = None) -> Tuple[bool, str]:
    """向Word文档添加表格

    Args:
        file_path: Word文档路径
        table_data: 表格数据（二维列表）
        workspace_path: 工作区路径

    Returns:
        Tuple[bool, str]: 是否成功和消息
    """
    try:
        # 处理文件路径
        if isinstance(file_path, str):
            file_path = Path(file_path)

        # 处理工作区路径
        if workspace_path is not None:
            workspace = Path(workspace_path)
        else:
            workspace = Path.cwd()

        # 处理相对路径
        if not file_path.is_absolute():
            file_path = workspace / file_path

        # 安全检查
        if not is_file_within_workspace(file_path, workspace):
            return False, "错误：文件不在当前工作区内"

        if not file_path.exists():
            return False, f"错误：文件不存在: {file_path}"

        # 打开文档并添加表格
        doc = Document(file_path)
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)

        # 填充表格数据
        for i, row in enumerate(table_data):
            for j, cell_value in enumerate(row):
                table.cell(i, j).text = cell_value

        doc.save(file_path)

        logger.info(f"成功向Word文档添加表格: {file_path}")
        return True, f"成功添加表格到文档"

    except Exception as e:
        return False, f"添加表格失败: {e}"


def search_in_word(file_path: Union[str, Path],
                   search_text: str,
                   workspace_path: Optional[Union[str, Path]] = None,
                   case_sensitive: bool = False) -> Tuple[bool, Dict]:
    """在Word文档中搜索文本

    Args:
        file_path: Word文档路径
        search_text: 要搜索的文本
        workspace_path: 工作区路径
        case_sensitive: 是否区分大小写

    Returns:
        Tuple[bool, Dict]: 是否成功和搜索结果
    """
    try:
        # 处理文件路径
        if isinstance(file_path, str):
            file_path = Path(file_path)

        # 处理工作区路径
        if workspace_path is not None:
            workspace = Path(workspace_path)
        else:
            workspace = Path.cwd()

        # 处理相对路径
        if not file_path.is_absolute():
            file_path = workspace / file_path

        # 安全检查
        if not is_file_within_workspace(file_path, workspace):
            return False, {"error": "文件不在当前工作区内"}

        if not file_path.exists():
            return False, {"error": f"文件不存在: {file_path}"}

        # 读取文档
        doc = Document(file_path)

        # 搜索文本
        results = []
        search_pattern = search_text if case_sensitive else search_text.lower()

        for para_idx, paragraph in enumerate(doc.paragraphs):
            text_to_search = paragraph.text if case_sensitive else paragraph.text.lower()

            if search_pattern in text_to_search:
                # 找到所有匹配位置
                start_pos = 0
                while True:
                    pos = text_to_search.find(search_pattern, start_pos)
                    if pos == -1:
                        break

                    results.append({
                        "paragraph_index": para_idx,
                        "paragraph_text": paragraph.text,
                        "position": pos,
                        "matched_text": paragraph.text[pos:pos + len(search_text)]
                    })
                    start_pos = pos + 1

        # 在表格中搜索
        table_results = []
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text if case_sensitive else cell.text.lower()
                    if search_pattern in cell_text:
                        table_results.append({
                            "table_index": table_idx,
                            "row": row_idx,
                            "column": cell_idx,
                            "cell_text": cell.text
                        })

        result = {
            "file_path": str(file_path),
            "search_text": search_text,
            "paragraph_matches": results,
            "table_matches": table_results,
            "total_matches": len(results) + len(table_results)
        }

        logger.info(f"在Word文档中找到 {len(results) + len(table_results)} 个匹配项")
        return True, result

    except Exception as e:
        logger.error(f"搜索Word文档失败: {e}")
        return False, {"error": f"搜索失败: {e}"}


if __name__ == '__main__':
    print('start..')
    status, msg = create_word_file('test.word', ['heading', 'para1', 'para2'],
                                   'D:\\app\pycharm_project\\remoteAgent\\workspace\\workspace_admin')
    print(msg)

    status, msg = read_word_file('test.docx', 'D:\\app\pycharm_project\\remoteAgent\\workspace\\workspace_admin')
    print(msg)
    status, msg = add_paragraph_to_word('test.docx', '新增的段落',
                                        'D:\\app\pycharm_project\\remoteAgent\\workspace\\workspace_admin')
    print(msg)
    print('end...')
